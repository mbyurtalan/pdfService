from flask import Flask, request, send_file, jsonify
from pypdf import PdfWriter
import tempfile
import os
import io
import re
import json
import base64
from docx import Document

app = Flask(__name__)


def leading_number(filename):
    match = re.match(r"^page_(\d+)", filename, re.IGNORECASE)
    return int(match.group(1)) if match else 999999999

@app.route('/', methods = ['GET'])
def status():
    return jsonify({'status': 'ok'}), 200

@app.route("/apply-regex", methods=["POST"])
def apply_regex():
    print("CONTENT TYPE:", request.content_type)
    print("HEADERS:", dict(request.headers))
    print("RAW LENGTH:", len(request.data))
    print("FIRST 200 BYTES:", request.data[:200])
    try:
        payload = request.get_json(force=True)
        doc_bytes = base64.b64decode(payload["$content"])
        doc = Document(io.BytesIO(doc_bytes))
        text_parts = []

        # Normal paragraphs
        for p in doc.paragraphs:
            text_parts.append(p.text)

        # Tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text_parts.append(cell.text)

        full_text = "\n".join(text_parts)
        match = re.search(
            r"SHEET\s*CONTD\.?\s*([A-Z])\s*(\d+)",
            full_text,
            re.IGNORECASE | re.DOTALL
        )

        if not match:
            return jsonify({
                "found": False,
                "page": None,
                "sheet": None
            })

        return jsonify({
            "found": True,
            "page": match.group(1).strip(),
            "sheet": int(match.group(2))
        })

    except Exception as e:
        print(e)
        return jsonify({"error": str(e)}), 500

@app.route("/merge-pdf", methods=["POST"])
def merge_pdf():

    try:
        files = request.get_json(force=True)

        # tolerate array of JSON strings
        normalized = []

        for item in files:
            if isinstance(item, str):
                item = json.loads(item)

            normalized.append(item)

        # sort by leading number
        normalized.sort(
            key=lambda x: leading_number(x["filename"])
        )

        with tempfile.TemporaryDirectory() as tmpdir:

            pdf_paths = []

            for item in normalized:

                pdf_path = os.path.join(
                    tmpdir,
                    os.path.basename(item["filename"])
                )

                with open(pdf_path, "wb") as f:
                    f.write(base64.b64decode(item["content"]))

                pdf_paths.append(pdf_path)

            merged_path = os.path.join(tmpdir, "merged.pdf")

            writer = PdfWriter()

            for pdf in pdf_paths:
                writer.append(pdf)

            with open(merged_path, "wb") as f:
                writer.write(f)

            return send_file(
                merged_path,
                mimetype="application/pdf",
                as_attachment=True,
                download_name="merged.pdf"
            )

    except Exception as e:
        return jsonify({"error": str(e)}), 500


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000)